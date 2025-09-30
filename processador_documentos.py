import os
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from PIL import Image

class ProcessadorDocumentos:
    def __init__(self, template_path, largura_imagem=6.0):
        self.template_path = template_path
        self.largura_imagem = largura_imagem
        self.qualidade_imagem = 1.0  # 1.0 = alta, 0.7 = média, 0.5 = baixa
        self.callback_progresso = None
        self.setup_logging()
    
    def setup_logging(self):
        # Usar logger centralizado do sistema principal
        import logging
        self.logger = logging.getLogger('IntegraEvidencias')
    
    def processar_diretorio(self, diretorio_raiz, arquivo_saida, callback_progresso=None):
        try:
            self.callback_progresso = callback_progresso
            self.logger.info(f"Iniciando processamento do diretório: {diretorio_raiz}")
            
            # Carregar template
            doc = Document(self.template_path)
            self.logger.info(f"Template carregado: {self.template_path}")
            
            # Obter estrutura de diretórios
            estrutura = self._obter_estrutura_diretorios(diretorio_raiz)
            
            # Inserir conteúdo após a primeira página
            self._inserir_conteudo(doc, estrutura)
            
            # Salvar documento
            doc.save(arquivo_saida)
            self.logger.info(f"Documento salvo: {arquivo_saida}")
            
            # Manter apenas formato .docx conforme solicitado
            # pdf_path = arquivo_saida.replace('.docx', '.pdf')
            # self._converter_para_pdf(arquivo_saida, pdf_path)
            
            return True
            
        except Exception as e:
            self.logger.error(f"Erro no processamento: {str(e)}")
            return False
    
    def _obter_estrutura_diretorios(self, diretorio_raiz):
        estrutura = []
        
        for root, dirs, files in os.walk(diretorio_raiz):
            # Filtrar apenas arquivos .jpg e .png
            imagens = [f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
            
            # Incluir todas as pastas na hierarquia, mesmo sem imagens diretas
            nivel = root.replace(diretorio_raiz, '').count(os.sep)
            nome_pasta = os.path.basename(root) if root != diretorio_raiz else "Evidências"
            caminho_relativo = root.replace(diretorio_raiz, '').strip(os.sep)
            
            estrutura.append({
                'nome': nome_pasta,
                'nivel': nivel,
                'caminho': root,
                'caminho_relativo': caminho_relativo,
                'imagens': [os.path.join(root, img) for img in sorted(imagens)]
            })
            
            self.logger.info(f"Pasta processada: {nome_pasta} (nível {nivel}) - {len(imagens)} imagens")
        
        return estrutura
    
    def _organizar_hierarquia(self, estrutura):
        # Ordenar por caminho para manter hierarquia correta
        estrutura_ordenada = sorted(estrutura, key=lambda x: (x['caminho_relativo'], x['nivel']))
        
        # Filtrar apenas pastas que têm imagens ou são pais de pastas com imagens
        estrutura_final = []
        
        for item in estrutura_ordenada:
            # Incluir se tem imagens próprias
            if item['imagens']:
                estrutura_final.append(item)
            # Ou se é pasta pai necessária para hierarquia
            else:
                # Verificar se alguma subpasta tem imagens
                tem_subpastas_com_imagens = any(
                    sub['caminho'].startswith(item['caminho'] + os.sep) and sub['imagens']
                    for sub in estrutura_ordenada
                )
                if tem_subpastas_com_imagens:
                    estrutura_final.append(item)
        
        return estrutura_final
    
    def _inserir_conteudo(self, doc, estrutura):
        # Organizar estrutura por hierarquia
        estrutura_organizada = self._organizar_hierarquia(estrutura)
        
        # Identificar e salvar TODOS os elementos da página 2 (TODOS OS DIREITOS RESERVADOS)
        paragrafos = list(doc.paragraphs)
        pagina2_elementos = []
        
        # Procurar por parágrafos que contenham "TODOS OS DIREITOS RESERVADOS"
        for i, p in enumerate(paragrafos):
            if "TODOS OS DIREITOS RESERVADOS" in p.text or "TOPAZ" in p.text:
                # Salvar todos os elementos da página 2 a partir deste ponto
                for j in range(i, len(paragrafos)):
                    if j < len(paragrafos) - 1:  # Não pegar o último se for diferente
                        pagina2_elementos.append(paragrafos[j]._element)
                        paragrafos[j]._element.getparent().remove(paragrafos[j]._element)
                break
        
        # Se não encontrou pelo texto, pegar da posição 1 em diante
        if not pagina2_elementos and len(paragrafos) > 1:
            for i in range(1, len(paragrafos)):
                pagina2_elementos.append(paragrafos[i]._element)
                paragrafos[i]._element.getparent().remove(paragrafos[i]._element)
        
        # Adicionar quebra de página após primeira página
        page_break = doc.add_paragraph()
        page_break.add_run().add_break(6)
        
        # Gerar conteúdo das evidências
        for item in estrutura_organizada:
            # Título
            titulo = doc.add_paragraph(item['nome'])
            if item['nivel'] == 0:
                titulo.style = 'Heading 1'
            elif item['nivel'] == 1:
                titulo.style = 'Heading 2'
            else:
                titulo.style = 'Heading 3'
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Imagens desta pasta
            for i, img_path in enumerate(item['imagens']):
                try:
                    # Otimizar imagem se necessário
                    img_path_otimizada = self._otimizar_imagem(img_path)
                    
                    # Parágrafo da imagem
                    img_para = doc.add_paragraph()
                    run = img_para.add_run()
                    
                    img_width, img_height = self._calcular_dimensoes_imagem(img_path_otimizada)
                    run.add_picture(img_path_otimizada, width=Inches(img_width), height=Inches(img_height))
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Nome da imagem
                    nome_para = doc.add_paragraph(f"Figura: {os.path.basename(img_path)}")
                    nome_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Espaçamento
                    doc.add_paragraph()
                    
                    self.logger.info(f"Imagem inserida: {os.path.basename(img_path)}")
                    
                    # Callback de progresso
                    if self.callback_progresso:
                        total_imagens = sum(len(pasta['imagens']) for pasta in estrutura_organizada)
                        imagem_atual = sum(len(pasta['imagens']) for pasta in estrutura_organizada[:estrutura_organizada.index(item)]) + i + 1
                        progresso = 0.2 + (0.7 * imagem_atual / total_imagens)
                        self.callback_progresso(progresso, f"Processando imagem {imagem_atual} de {total_imagens}...")
                    
                    # Limpar arquivo temporário se foi criado
                    if img_path_otimizada != img_path and os.path.exists(img_path_otimizada):
                        try:
                            os.remove(img_path_otimizada)
                        except:
                            pass
                    
                except Exception as e:
                    self.logger.error(f"Erro ao inserir imagem {img_path}: {str(e)}")
        
        # Forçar quebra de página antes da página final
        page_break_final = doc.add_paragraph()
        page_break_final.add_run().add_break(6)
        
        # Restaurar página 2 (TODOS OS DIREITOS RESERVADOS) no final
        for elemento in pagina2_elementos:
            doc._body._element.append(elemento)
        
        self.logger.info("Página 'TODOS OS DIREITOS RESERVADOS' movida para o final do documento")
    

    
    def _otimizar_imagem(self, img_path):
        """Otimiza imagem baseada na qualidade selecionada"""
        if self.qualidade_imagem >= 1.0:
            return img_path  # Sem otimização
        
        try:
            with Image.open(img_path) as img:
                # Converter para RGB se necessário
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                # Redimensionar se muito grande
                max_size = int(2000 * self.qualidade_imagem)
                if max(img.size) > max_size:
                    img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                
                # Salvar com compressão
                temp_path = img_path + '_temp.jpg'
                quality = int(85 * self.qualidade_imagem)
                img.save(temp_path, 'JPEG', quality=quality, optimize=True)
                
                return temp_path
        except Exception as e:
            self.logger.error(f"Erro ao otimizar imagem {img_path}: {str(e)}")
            return img_path
    
    def _calcular_dimensoes_imagem(self, img_path):
        try:
            with Image.open(img_path) as img:
                width, height = img.size
                aspect_ratio = width / height
                
                # Definir dimensões máximas (em polegadas)
                max_width = min(self.largura_imagem, 6.5)  # Margem de segurança
                max_height = 8.0  # Altura máxima para evitar cortes
                
                # Calcular dimensões mantendo proporção
                if width > height:  # Paisagem
                    img_width = max_width
                    img_height = max_width / aspect_ratio
                    if img_height > max_height:
                        img_height = max_height
                        img_width = max_height * aspect_ratio
                else:  # Retrato
                    img_height = min(max_height, max_width / aspect_ratio)
                    img_width = img_height * aspect_ratio
                    if img_width > max_width:
                        img_width = max_width
                        img_height = max_width / aspect_ratio
                
                return img_width, img_height
        except:
            return self.largura_imagem, self.largura_imagem * 0.75
    
    def _converter_para_pdf(self, docx_path, pdf_path):
        try:
            # Tentar conversão usando win32com (Windows)
            import win32com.client
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            
            self.logger.info(f"PDF gerado: {pdf_path}")
            
        except ImportError:
            self.logger.warning("win32com não disponível. PDF não foi gerado.")
        except Exception as e:
            self.logger.error(f"Erro na conversão para PDF: {str(e)}")